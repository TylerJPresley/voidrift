"""Document format extraction tool (REQ-U-19).

Soft dependencies: pymupdf (PDF), python-docx (DOCX), openpyxl (XLSX).
Returns plaintext/markdown. Missing libraries produce install hints.
"""

from __future__ import annotations

from pathlib import Path

SUPPORTED = {".pdf", ".docx", ".xlsx"}


def read_document(path: str, project_root: str) -> str:
    """Extract text from a binary document file.

    Args:
        path: File path relative to project root.
        project_root: Absolute path to the project directory.

    Returns:
        Extracted text as plaintext/markdown, or an error string.
    """
    root = Path(project_root).resolve()
    resolved = (root / path).resolve()
    try:
        resolved.relative_to(root)
    except ValueError:
        return f"Error: path '{path}' is outside the project root."

    if not resolved.exists():
        return f"Error: file not found: {path}"

    ext = resolved.suffix.lower()
    if ext not in SUPPORTED:
        return f"Error: unsupported format '{ext}'. Supported: {', '.join(sorted(SUPPORTED))}"

    if ext == ".pdf":
        return _read_pdf(resolved)
    elif ext == ".docx":
        return _read_docx(resolved)
    else:
        return _read_xlsx(resolved)


def _read_pdf(path: Path) -> str:
    try:
        import fitz  # pymupdf
    except ImportError:
        return "Error: pymupdf is not installed. Run: pip install pymupdf"
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()
    if not pages:
        return "(empty document)"
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return "Error: python-docx is not installed. Run: pip install python-docx"
    doc = docx.Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("Heading"):
            lines.append(f"#### {text}")
        else:
            lines.append(text)
    for table in doc.tables:
        lines.append("")
        lines.append(_table_to_markdown(
            [[cell.text.strip() for cell in row.cells] for row in table.rows]
        ))
    return "\n\n".join(lines) if lines else "(empty document)"


def _read_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        return "Error: openpyxl is not installed. Run: pip install openpyxl"
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    sections: list[str] = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = [[str(cell) if cell is not None else "" for cell in row] for row in ws.values]
        if not rows:
            continue
        sections.append(f"## {sheet}\n\n{_table_to_markdown(rows)}")
    wb.close()
    return "\n\n".join(sections) if sections else "(empty document)"


def _table_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header = rows[0]
    sep = ["---"] * len(header)
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(sep) + " |"]
    for row in rows[1:]:
        # Pad row to header length
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(padded[:len(header)]) + " |")
    return "\n".join(lines)
